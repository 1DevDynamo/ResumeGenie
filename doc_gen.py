import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# ================= FONT =================
def set_font(run, size=11, bold=False, color=(0,0,0)):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(*color)

# ================= PAGE MARGINS =================
def set_margins(doc):

    section = doc.sections[0]

    section.top_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    section.bottom_margin = Inches(1.0)
    section.footer_distance = Inches(0.3)

# ================= SECTION DIVIDER =================
def add_horizontal_line(doc):
    p = doc.add_paragraph()
    fmt = p.paragraph_format

    fmt.space_before = Pt(2)
    fmt.space_after = Pt(6)
    fmt.keep_with_next = True 

    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')

    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'B4B4B4')

    borders.append(bottom)
    p_pr.append(borders)

# ================= SECTION TITLE =================
def add_section_title(doc, title):
    add_horizontal_line(doc)

    p = doc.add_paragraph()
    fmt = p.paragraph_format

    fmt.space_before = Pt(12)
    fmt.space_after = Pt(6)
    fmt.line_spacing = 1.0
    fmt.keep_with_next = True   # 🔥 prevents heading alone

    run = p.add_run(title)
    set_font(run, 13, bold=True, color=(0,102,204))

# ================= LEFT-RIGHT ALIGN =================
def add_left_right_line(doc, left_text, right_text, bold_left=False):
    p = doc.add_paragraph()
    fmt = p.paragraph_format

    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.08

    tab_stops = fmt.tab_stops
    tab_stops.add_tab_stop(Inches(6.4), WD_TAB_ALIGNMENT.RIGHT)

    left_run = p.add_run(left_text)
    set_font(left_run, 11, bold_left)

    p.add_run("\t")

    right_run = p.add_run(right_text)
    set_font(right_run, 10.5, color=(100,100,100))

# ================= CUSTOM BULLET =================
def add_bullet(doc, text):
    p = doc.add_paragraph()
    fmt = p.paragraph_format

    fmt.left_indent = Inches(0.25)
    fmt.first_line_indent = Inches(-0.25)
    fmt.space_after = Pt(3)
    fmt.line_spacing = 1.08
    fmt.keep_together = True    # 🔥 prevents bullet split

    bullet_run = p.add_run("• ")
    set_font(bullet_run)

    pattern = r"\*\*(.*?)\*\*"
    last_end = 0

    for match in re.finditer(pattern, text):
        start, end = match.span()

        if start > last_end:
            normal = text[last_end:start]
            run = p.add_run(normal)
            set_font(run)

        bold = match.group(1)
        run = p.add_run(bold)
        set_font(run, bold=True)

        last_end = end

    if last_end < len(text):
        run = p.add_run(text[last_end:])
        set_font(run)

# ================= PROJECT TITLE =================
def add_project_title(doc, title):
    p = doc.add_paragraph()
    fmt = p.paragraph_format

    fmt.space_after = Pt(0)
    fmt.keep_with_next = True   # 🔥 title sticks to bullets

    run = p.add_run(title)
    set_font(run, 11, bold=True)

# ================= SUMMARY BLOCK =================
def add_summary_block(doc, text):
    p = doc.add_paragraph()
    fmt = p.paragraph_format

    fmt.space_before = Pt(4)
    fmt.space_after = Pt(8)
    fmt.line_spacing = 1.08

    pattern = r"\*\*(.*?)\*\*"
    last_end = 0

    for match in re.finditer(pattern, text):
        start, end = match.span()

        if start > last_end:
            normal = text[last_end:start]
            run = p.add_run(normal)
            set_font(run, 10.5)

        bold = match.group(1)
        run = p.add_run(bold)
        set_font(run, 10.5, bold=True)

        last_end = end

    if last_end < len(text):
        run = p.add_run(text[last_end:])
        set_font(run, 10.5)

# ================= HEADER =================
def add_header(doc, header):

    name_para = doc.add_paragraph()
    name_run = name_para.add_run(header.get("name", ""))
    set_font(name_run, 20, bold=True)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def icon_run(p, icon, text):
        r1 = p.add_run(icon + " ")
        set_font(r1, 10)
        r2 = p.add_run(text + "   ")
        set_font(r2, 10, color=(0,102,204))

    if header.get("phone"):
        icon_run(contact_para, "☎", header.get("phone"))
    if header.get("email"):
        icon_run(contact_para, "✉", header.get("email"))
    if header.get("linkedin"):
        icon_run(contact_para, "🔗", header.get("linkedin"))
    if header.get("github"):
        icon_run(contact_para, "🐙", header.get("github"))

# ================= MAIN GENERATOR =================
def generate_docx_from_template(resume):

    doc = Document()
    set_margins(doc)

    add_header(doc, resume.get("header", {}))

    summary = resume.get("summary", "").strip()
    if summary:
        add_section_title(doc, "Summary")
        add_summary_block(doc, summary)

    education = resume.get("education", [])
    if education:
        add_section_title(doc, "Education")
        for edu in education:
            add_left_right_line(doc, edu.get("degree", ""), edu.get("duration", ""), True)
            add_left_right_line(doc, edu.get("institution", ""), edu.get("grade", ""))

    experience = resume.get("experience", [])
    if experience:
        add_section_title(doc, "Experience")
        for exp in experience:
            role_company = f"{exp.get('role','')} — {exp.get('company','')}"
            add_left_right_line(doc, role_company, exp.get("duration", ""), True)
            for bullet in exp.get("bullets", []):
                add_bullet(doc, bullet)

    projects = resume.get("projects", [])
    if projects:
        add_section_title(doc, "Projects")
        for proj in projects:
            add_project_title(doc, proj.get("title", ""))
            for bullet in proj.get("bullets", []):
                add_bullet(doc, bullet)

    coursework = resume.get("coursework", [])
    coursework = resume.get("coursework", [])
    if coursework:
        add_section_title(doc, "Relevant Coursework")

        p = doc.add_paragraph()
        fmt = p.paragraph_format

        fmt.line_spacing = 1.08
        fmt.keep_together = True
        # no keep_with_next needed (single line)

        run = p.add_run(" • ".join(coursework))
        set_font(run, 10.5, bold=False)

    skills = resume.get("skills", {})
    if skills:
        add_section_title(doc, "Technical Skills")

        skill_items = list(skills.items())

        for i, (category, items) in enumerate(skill_items):

            p = doc.add_paragraph()
            fmt = p.paragraph_format

            fmt.line_spacing = 1.08
            fmt.keep_together = True

            # 🔥 glue all lines except last one
            if i != len(skill_items) - 1:
                fmt.keep_with_next = True

            run = p.add_run(f"{category}: {', '.join(items)}")
            set_font(run, 10.5, bold=False)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output = os.path.join(BASE_DIR, f"Generated_Resume_{timestamp}.docx")

    doc.save(output)
    return output